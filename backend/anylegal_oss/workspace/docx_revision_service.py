"""
DOCX Revision Service - Native Word Track Changes

Provides native Word track changes (revision markup) using the docx-revisions library.
Creates OOXML revision markup that Word recognizes natively.

This is part of the Hybrid DOCX Architecture for producing professional
redlined documents that can be opened directly in Microsoft Word.

Dependencies:
- docx-revisions: Native Word track changes support
- python-docx: DOCX manipulation
"""

import io
import logging
import re
from typing import Optional, Tuple, Dict, Any, List
from datetime import datetime

logger = logging.getLogger(__name__)

_docx_revisions = None
_Document = None

def _ensure_docx_revisions():
    """Lazy load docx-revisions for track changes."""
    global _docx_revisions
    if _docx_revisions is None:
        try:
            import docx_revisions
            _docx_revisions = docx_revisions
        except ImportError:
            logger.error("docx-revisions not installed. Run: pip install docx-revisions")
            raise ImportError("docx-revisions library required for track changes")
    return _docx_revisions

def _ensure_docx():
    """Lazy load python-docx."""
    global _Document
    if _Document is None:
        try:
            from docx import Document
            _Document = Document
        except ImportError:
            logger.error("python-docx not installed")
            raise ImportError("python-docx library required")
    return _Document

class DocxRevisionService:
    """
    Service for creating and managing Word track changes.

    Uses docx-revisions to create native OOXML revision markup that
    Word displays as proper tracked changes.
    """

    DEFAULT_AUTHOR = "AnyLegal"

    @classmethod
    def create_redlined_docx(
        cls,
        original_docx: bytes,
        revised_text: str,
        author: str = None,
        change_description: str = None
    ) -> Tuple[bytes, Dict[str, Any]]:
        """
        Create a redlined DOCX with track changes from original and revised text.

        Args:
            original_docx: Original DOCX file bytes
            revised_text: The revised plain text content
            author: Author name for revision attribution
            change_description: Optional description of changes

        Returns:
            Tuple of (redlined_docx_bytes, stats_dict)
        """
        docx_revisions = _ensure_docx_revisions()
        Document = _ensure_docx()

        author = author or cls.DEFAULT_AUTHOR

        try:

            doc = Document(io.BytesIO(original_docx))

            original_text = '\n\n'.join(p.text for p in doc.paragraphs)

            redlined_doc, stats = cls._apply_revisions(
                doc, original_text, revised_text, author
            )

            output = io.BytesIO()
            redlined_doc.save(output)
            output.seek(0)

            return output.read(), stats

        except Exception as e:
            logger.error(f"Failed to create redlined DOCX: {e}")
            raise ValueError(f"Track changes creation failed: {str(e)}")

    @classmethod
    def _apply_revisions(
        cls,
        doc,
        original_text: str,
        revised_text: str,
        author: str
    ) -> Tuple[Any, Dict[str, Any]]:
        """
        Apply track changes to document using docx-revisions.

        This implementation uses a paragraph-by-paragraph comparison
        to create tracked changes.
        """
        docx_revisions = _ensure_docx_revisions()

        stats = {
            "insertions": 0,
            "deletions": 0,
            "unchanged": 0,
            "paragraphs_modified": 0
        }

        original_paras = original_text.split('\n\n')
        revised_paras = revised_text.split('\n\n')

        import difflib
        matcher = difflib.SequenceMatcher(None, original_paras, revised_paras)

        modifications = []

        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'equal':
                stats["unchanged"] += (i2 - i1)
            elif tag == 'replace':

                for idx in range(i1, i2):
                    if idx < len(doc.paragraphs) and (j1 + idx - i1) < len(revised_paras):
                        modifications.append({
                            "para_idx": idx,
                            "old_text": original_paras[idx] if idx < len(original_paras) else "",
                            "new_text": revised_paras[j1 + idx - i1],
                            "type": "replace"
                        })
                        stats["paragraphs_modified"] += 1
            elif tag == 'delete':
                for idx in range(i1, i2):
                    modifications.append({
                        "para_idx": idx,
                        "old_text": original_paras[idx] if idx < len(original_paras) else "",
                        "new_text": "",
                        "type": "delete"
                    })
                    stats["deletions"] += 1
            elif tag == 'insert':
                for jdx in range(j1, j2):
                    modifications.append({
                        "para_idx": i1,                              
                        "old_text": "",
                        "new_text": revised_paras[jdx],
                        "type": "insert"
                    })
                    stats["insertions"] += 1

        try:

            if hasattr(docx_revisions, 'add_revision'):
                for mod in modifications:
                    if mod["para_idx"] < len(doc.paragraphs):
                        para = doc.paragraphs[mod["para_idx"]]
                        if mod["type"] == "replace":

                            docx_revisions.add_revision(
                                para,
                                mod["old_text"],
                                mod["new_text"],
                                author=author,
                                date=datetime.now()
                            )
                        elif mod["type"] == "delete":

                            docx_revisions.add_deletion(
                                para,
                                mod["old_text"],
                                author=author,
                                date=datetime.now()
                            )
                        elif mod["type"] == "insert":

                            docx_revisions.add_insertion(
                                para,
                                mod["new_text"],
                                author=author,
                                date=datetime.now()
                            )
            else:

                cls._apply_manual_revisions(doc, modifications, author)

        except Exception as e:
            logger.warning(f"docx-revisions API call failed, using fallback: {e}")
            cls._apply_manual_revisions(doc, modifications, author)

        return doc, stats

    @classmethod
    def _apply_manual_revisions(
        cls,
        doc,
        modifications: List[Dict[str, Any]],
        author: str
    ) -> None:
        """
        Manually apply OOXML revision markup when docx-revisions API is unavailable.

        This creates basic track changes using direct XML manipulation.
        """
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from lxml import etree

        revision_id = 1
        date_str = datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ')

        for mod in modifications:
            if mod["para_idx"] >= len(doc.paragraphs):
                continue

            para = doc.paragraphs[mod["para_idx"]]

            if mod["type"] == "replace" and mod["old_text"] and mod["new_text"]:

                for run in para.runs:

                    del_elem = OxmlElement('w:del')
                    del_elem.set(qn('w:id'), str(revision_id))
                    del_elem.set(qn('w:author'), author)
                    del_elem.set(qn('w:date'), date_str)

                    run._r.getparent().replace(run._r, del_elem)
                    del_elem.append(run._r)
                    revision_id += 1

                ins_elem = OxmlElement('w:ins')
                ins_elem.set(qn('w:id'), str(revision_id))
                ins_elem.set(qn('w:author'), author)
                ins_elem.set(qn('w:date'), date_str)

                new_run = para.add_run(mod["new_text"])
                new_run._r.getparent().replace(new_run._r, ins_elem)
                ins_elem.append(new_run._r)
                revision_id += 1

            elif mod["type"] == "delete" and mod["old_text"]:

                for run in para.runs:
                    del_elem = OxmlElement('w:del')
                    del_elem.set(qn('w:id'), str(revision_id))
                    del_elem.set(qn('w:author'), author)
                    del_elem.set(qn('w:date'), date_str)
                    run._r.getparent().replace(run._r, del_elem)
                    del_elem.append(run._r)
                    revision_id += 1

            elif mod["type"] == "insert" and mod["new_text"]:

                ins_elem = OxmlElement('w:ins')
                ins_elem.set(qn('w:id'), str(revision_id))
                ins_elem.set(qn('w:author'), author)
                ins_elem.set(qn('w:date'), date_str)

                new_run = para.add_run(mod["new_text"])
                new_run._r.getparent().replace(new_run._r, ins_elem)
                ins_elem.append(new_run._r)
                revision_id += 1

    @classmethod
    def accept_all_revisions(cls, docx_bytes: bytes) -> bytes:
        """
        Accept all tracked changes in a DOCX document.

        Args:
            docx_bytes: DOCX file with tracked changes

        Returns:
            Clean DOCX with all changes accepted
        """
        docx_revisions = _ensure_docx_revisions()
        Document = _ensure_docx()

        try:
            doc = Document(io.BytesIO(docx_bytes))

            if hasattr(docx_revisions, 'accept_all'):
                docx_revisions.accept_all(doc)
            else:

                cls._accept_all_manual(doc)

            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            return output.read()

        except Exception as e:
            logger.error(f"Failed to accept revisions: {e}")
            raise ValueError(f"Accept revisions failed: {str(e)}")

    @classmethod
    def _accept_all_manual(cls, doc) -> None:
        """Manually accept all revisions by removing markup."""
        from docx.oxml.ns import qn

        for body_elem in doc.element.body:

            for del_elem in body_elem.findall('.//' + qn('w:del'), body_elem.nsmap):
                parent = del_elem.getparent()
                parent.remove(del_elem)

            for ins_elem in body_elem.findall('.//' + qn('w:ins'), body_elem.nsmap):
                parent = ins_elem.getparent()
                index = list(parent).index(ins_elem)
                for child in ins_elem:
                    parent.insert(index, child)
                    index += 1
                parent.remove(ins_elem)

    @classmethod
    def reject_all_revisions(cls, docx_bytes: bytes) -> bytes:
        """
        Reject all tracked changes in a DOCX document.

        Args:
            docx_bytes: DOCX file with tracked changes

        Returns:
            Original DOCX with all changes rejected
        """
        docx_revisions = _ensure_docx_revisions()
        Document = _ensure_docx()

        try:
            doc = Document(io.BytesIO(docx_bytes))

            if hasattr(docx_revisions, 'reject_all'):
                docx_revisions.reject_all(doc)
            else:
                cls._reject_all_manual(doc)

            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            return output.read()

        except Exception as e:
            logger.error(f"Failed to reject revisions: {e}")
            raise ValueError(f"Reject revisions failed: {str(e)}")

    @classmethod
    def _reject_all_manual(cls, doc) -> None:
        """Manually reject all revisions."""
        from docx.oxml.ns import qn

        for body_elem in doc.element.body:

            for del_elem in body_elem.findall('.//' + qn('w:del'), body_elem.nsmap):
                parent = del_elem.getparent()
                index = list(parent).index(del_elem)
                for child in del_elem:
                    parent.insert(index, child)
                    index += 1
                parent.remove(del_elem)

            for ins_elem in body_elem.findall('.//' + qn('w:ins'), body_elem.nsmap):
                parent = ins_elem.getparent()
                parent.remove(ins_elem)

    @classmethod
    def get_revision_stats(cls, docx_bytes: bytes, with_snippets: bool = False) -> Dict[str, Any]:
        """
        Get statistics about tracked changes in a document.

        Args:
            docx_bytes: Document bytes
            with_snippets: When True, also return per-revision detail —
                ``revisions`` list with ``id``, ``type``, ``author``,
                ``date``, ``text_snippet`` (≤80 chars), ``context_around``
                (≤200 chars). Use this when the model needs to pick
                specific revision IDs to accept/reject. Default False
                keeps existing call shape cheap.

        Returns:
            Dict with revision counts, authors, IDs (and per-revision
            detail when with_snippets=True).
        """
        Document = _ensure_docx()

        try:
            doc = Document(io.BytesIO(docx_bytes))

            from docx.oxml.ns import qn

            stats: Dict[str, Any] = {
                "insertions": 0,
                "deletions": 0,
                "authors": set(),
                "has_revisions": False,
                "revision_ids": [],
            }
            if with_snippets:
                stats["revisions"] = []

            ID_ATTR = qn('w:id')
            AUTHOR_ATTR = qn('w:author')
            DATE_ATTR = qn('w:date')

            def _extract_text(elem) -> str:
                """Get text from all <w:t> / <w:delText> descendants."""
                t_tags = (qn('w:t'), qn('w:delText'))
                parts = []
                for child in elem.iter():
                    if child.tag in t_tags and child.text:
                        parts.append(child.text)
                return "".join(parts)

            def _context_around(elem, before_chars: int = 100, after_chars: int = 100) -> str:
                """Get ±chars of text around the element within its paragraph."""

                p = elem
                while p is not None and p.tag != qn('w:p'):
                    p = p.getparent()
                if p is None:
                    return ""
                full = _extract_text(p)
                target = _extract_text(elem)
                if not target:
                    return full[: before_chars + after_chars]
                idx = full.find(target)
                if idx < 0:
                    return full[: before_chars + after_chars]
                start = max(0, idx - before_chars)
                end = min(len(full), idx + len(target) + after_chars)
                return full[start:end]

            for body_elem in doc.element.body:
                ins_elems = body_elem.findall('.//' + qn('w:ins'))
                del_elems = body_elem.findall('.//' + qn('w:del'))

                stats["insertions"] += len(ins_elems)
                stats["deletions"] += len(del_elems)

                for elem in ins_elems + del_elems:
                    author = elem.get(AUTHOR_ATTR)
                    if author:
                        stats["authors"].add(author)
                    rid = elem.get(ID_ATTR)
                    if rid is not None:
                        try:
                            stats["revision_ids"].append(int(rid))
                        except ValueError:
                            pass

                if with_snippets:
                    for elem in ins_elems:
                        text = _extract_text(elem)
                        rid = elem.get(ID_ATTR)
                        stats["revisions"].append({
                            "id": int(rid) if rid is not None else None,
                            "type": "insertion",
                            "author": elem.get(AUTHOR_ATTR),
                            "date": elem.get(DATE_ATTR),
                            "text_snippet": text[:80],
                            "context_around": _context_around(elem),
                        })
                    for elem in del_elems:
                        text = _extract_text(elem)
                        rid = elem.get(ID_ATTR)
                        stats["revisions"].append({
                            "id": int(rid) if rid is not None else None,
                            "type": "deletion",
                            "author": elem.get(AUTHOR_ATTR),
                            "date": elem.get(DATE_ATTR),
                            "text_snippet": text[:80],
                            "context_around": _context_around(elem),
                        })

            stats["has_revisions"] = stats["insertions"] > 0 or stats["deletions"] > 0
            stats["authors"] = list(stats["authors"])
            stats["total_changes"] = stats["insertions"] + stats["deletions"]
            stats["revision_ids"] = sorted(set(stats["revision_ids"]))

            if with_snippets:
                stats["revisions"].sort(key=lambda r: r["id"] if r["id"] is not None else 0)

            return stats

        except Exception as e:
            logger.error(f"Failed to get revision stats: {e}")
            return {"error": str(e)}

def accept_all_changes(docx_bytes: bytes) -> bytes:
    """Accept all tracked changes in a document."""
    return DocxRevisionService.accept_all_revisions(docx_bytes)

def reject_all_changes(docx_bytes: bytes) -> bytes:
    """Reject all tracked changes in a document."""
    return DocxRevisionService.reject_all_revisions(docx_bytes)
