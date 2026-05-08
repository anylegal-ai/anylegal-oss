"""
DOCX Service - Hybrid DOCX Architecture

Provides DOCX to HTML conversion and native Word document manipulation.
Part of the Hybrid DOCX Architecture where DOCX is the source of truth
and HTML is used for Tiptap display/editing.

Dependencies:
- python-docx: Create and modify DOCX files
- mammoth: DOCX to HTML conversion (high fidelity)
- docx-revisions: Native Word track changes (Phase 3)
"""

import io
import logging
import re
from typing import Optional, Tuple, Dict, Any, List
from datetime import datetime

logger = logging.getLogger(__name__)

_mammoth = None
_Document = None

def _ensure_mammoth():
    """Lazy load mammoth for DOCX to HTML conversion."""
    global _mammoth
    if _mammoth is None:
        try:
            import mammoth
            _mammoth = mammoth
            logger.info("mammoth loaded successfully from: %s", mammoth.__file__)
        except ImportError:
            logger.error("mammoth not installed. Run: pip install mammoth")
            raise ImportError("mammoth library required for DOCX conversion")
    return _mammoth

def _ensure_docx():
    """Lazy load python-docx for DOCX manipulation."""
    global _Document
    if _Document is None:
        try:
            from docx import Document
            _Document = Document
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            raise ImportError("python-docx library required for DOCX manipulation")
    return _Document

class DocxService:
    """
    Service for DOCX document operations in the Hybrid Architecture.

    Handles:
    - DOCX → HTML conversion for Tiptap display
    - HTML → DOCX conversion for export
    - Basic DOCX manipulation (read, create, modify)
    """

    STYLE_MAP = """
        p[style-name='Heading 1'] => h1:fresh
        p[style-name='Heading 2'] => h2:fresh
        p[style-name='Heading 3'] => h3:fresh
        p[style-name='Heading 4'] => h4:fresh
        p[style-name='Title'] => h1.document-title:fresh
        b => strong
        i => em
        u => u
        strike => del
    """

    @classmethod
    def docx_to_html(cls, docx_bytes: bytes) -> Tuple[str, List[Dict[str, Any]]]:
        """
        Convert DOCX bytes to HTML for Tiptap display.

        Args:
            docx_bytes: Raw DOCX file bytes

        Returns:
            Tuple of (html_content, messages) where messages contains warnings/info
        """
        mammoth = _ensure_mammoth()

        try:
            docx_stream = io.BytesIO(docx_bytes)

            result = mammoth.convert_to_html(
                docx_stream,
                style_map=cls.STYLE_MAP
            )

            html = result.value
            messages = [
                {"type": m.type, "message": m.message}
                for m in result.messages
            ]

            html = cls._postprocess_html_for_tiptap(html)

            logger.info(f"DOCX to HTML conversion complete: {len(html)} chars, {len(messages)} messages")
            return html, messages

        except Exception as e:
            logger.error(f"DOCX to HTML conversion failed: {e}")
            raise ValueError(f"Failed to convert DOCX: {str(e)}")

    @classmethod
    def _postprocess_html_for_tiptap(cls, html: str) -> str:
        """
        Post-process mammoth HTML output for better Tiptap compatibility.

        - Ensure proper paragraph tags
        - Clean up empty elements
        - Normalize whitespace
        - Mark DOCX revision tags for Tiptap track changes extension
        """

        html = re.sub(r'<p>\s*</p>', '', html)

        html = re.sub(r'<blockquote>\s*([^<])', r'<blockquote><p>\1', html)
        html = re.sub(r'([^>])\s*</blockquote>', r'\1</p></blockquote>', html)

        html = re.sub(r'(<br\s*/?>){3,}', '<br><br>', html)

        html = re.sub(
            r'<ins([^>]*)>',
            r'<ins\1 data-docx-revision="true">',
            html
        )
        html = re.sub(
            r'<del([^>]*)>',
            r'<del\1 data-docx-revision="true">',
            html
        )

        return html.strip()

    @classmethod
    def html_to_docx(
        cls,
        html_content: str,
        title: Optional[str] = None,
        base_docx: Optional[bytes] = None
    ) -> bytes:
        """
        Convert HTML content to DOCX format.

        This is a basic conversion that creates a new DOCX from HTML.
        For preserving original DOCX formatting, use update_docx_content() instead.

        Args:
            html_content: HTML string to convert
            title: Optional document title
            base_docx: Optional base DOCX to use as template (preserves styles)

        Returns:
            DOCX file as bytes
        """
        Document = _ensure_docx()

        try:
            if base_docx:

                doc = Document(io.BytesIO(base_docx))

                for element in doc.element.body[:]:
                    doc.element.body.remove(element)
            else:

                doc = Document()

            cls._html_to_docx_content(doc, html_content, title)

            output = io.BytesIO()
            doc.save(output)
            output.seek(0)

            docx_bytes = output.read()
            logger.info(f"HTML to DOCX conversion complete: {len(docx_bytes)} bytes")
            return docx_bytes

        except Exception as e:
            logger.error(f"HTML to DOCX conversion failed: {e}")
            raise ValueError(f"Failed to create DOCX: {str(e)}")

    @classmethod
    def _html_to_docx_content(
        cls,
        doc,
        html_content: str,
        title: Optional[str] = None
    ) -> None:
        """
        Parse HTML and add content to a python-docx Document.

        This is a simplified converter that handles common elements.
        For complex documents, the original DOCX should be preserved.
        """
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        if title:
            title_para = doc.add_heading(title, level=0)

        blocks = re.split(r'</(?:p|h[1-6]|blockquote|ul|ol)>', html_content)

        for block in blocks:
            block = block.strip()
            if not block:
                continue

            if block.startswith('<h1'):
                text = cls._strip_html_tags(block)
                doc.add_heading(text, level=1)
            elif block.startswith('<h2'):
                text = cls._strip_html_tags(block)
                doc.add_heading(text, level=2)
            elif block.startswith('<h3'):
                text = cls._strip_html_tags(block)
                doc.add_heading(text, level=3)
            elif block.startswith('<h4'):
                text = cls._strip_html_tags(block)
                doc.add_heading(text, level=4)
            elif block.startswith('<blockquote'):
                text = cls._strip_html_tags(block)
                para = doc.add_paragraph(text)
                para.paragraph_format.left_indent = Inches(0.5)
            elif block.startswith('<li'):
                text = cls._strip_html_tags(block)
                doc.add_paragraph(text, style='List Bullet')
            else:

                text = cls._strip_html_tags(block)
                if text:
                    para = doc.add_paragraph()
                    cls._add_formatted_text(para, block)

    @classmethod
    def _strip_html_tags(cls, html: str) -> str:
        """Remove HTML tags from a string."""

        text = re.sub(r'<[^>]+>', '', html)

        text = text.replace('&nbsp;', ' ')
        text = text.replace('&amp;', '&')
        text = text.replace('&lt;', '<')
        text = text.replace('&gt;', '>')
        text = text.replace('&quot;', '"')
        return text.strip()

    @classmethod
    def _add_formatted_text(cls, paragraph, html: str) -> None:
        """
        Add formatted text to a paragraph, preserving bold/italic/underline.
        """

        html = re.sub(r'^<[^>]+>', '', html)

        parts = re.split(r'(</?(?:strong|b|em|i|u|del)>)', html)

        bold = False
        italic = False
        underline = False
        strikethrough = False

        for part in parts:
            if part in ['<strong>', '<b>']:
                bold = True
            elif part in ['</strong>', '</b>']:
                bold = False
            elif part in ['<em>', '<i>']:
                italic = True
            elif part in ['</em>', '</i>']:
                italic = False
            elif part == '<u>':
                underline = True
            elif part == '</u>':
                underline = False
            elif part == '<del>':
                strikethrough = True
            elif part == '</del>':
                strikethrough = False
            elif part.strip():

                text = re.sub(r'<[^>]+>', '', part)
                text = text.replace('&nbsp;', ' ').replace('&amp;', '&')
                text = text.replace('&lt;', '<').replace('&gt;', '>').replace('&quot;', '"')
                if text:
                    run = paragraph.add_run(text)
                    run.bold = bold
                    run.italic = italic
                    run.underline = underline
                    run.font.strike = strikethrough

    @classmethod
    def read_docx_text(cls, docx_bytes: bytes) -> str:
        """
        Extract plain text from a DOCX file.

        Args:
            docx_bytes: Raw DOCX file bytes

        Returns:
            Plain text content
        """
        Document = _ensure_docx()

        try:
            doc = Document(io.BytesIO(docx_bytes))

            paragraphs = []
            for para in doc.paragraphs:
                paragraphs.append(para.text)

            return '\n\n'.join(paragraphs)

        except Exception as e:
            logger.error(f"Failed to read DOCX text: {e}")
            raise ValueError(f"Failed to read DOCX: {str(e)}")

    @classmethod
    def create_docx_from_text(
        cls,
        text: str,
        title: Optional[str] = None
    ) -> bytes:
        """
        Create a simple DOCX from plain text.

        Args:
            text: Plain text content
            title: Optional document title

        Returns:
            DOCX file as bytes
        """
        Document = _ensure_docx()

        doc = Document()

        if title:
            doc.add_heading(title, level=0)

        paragraphs = text.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                doc.add_paragraph(para_text.strip())

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        return output.read()

    @classmethod
    def _create_legal_template(cls):
        """Create a base DOCX with professional legal document styles.

        Sets Times New Roman as the theme font (both major and minor) so
        that any run referencing the theme inherits the right typeface.
        Also configures heading/body styles and page margins.
        """
        from docx.shared import Pt, RGBColor, Inches
        from lxml import etree

        Document = _ensure_docx()
        doc = Document()

        ns = {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}
        for rel in doc.part.rels.values():
            if 'theme' in rel.reltype:
                theme_xml = etree.fromstring(rel.target_part.blob)
                for tag in ('majorFont', 'minorFont'):
                    font_el = theme_xml.find(f'.//a:{tag}', ns)
                    if font_el is not None:
                        latin = font_el.find('a:latin', ns)
                        if latin is not None:
                            latin.set('typeface', 'Times New Roman')
                rel.target_part._blob = etree.tostring(theme_xml, xml_declaration=True,
                                                        encoding='UTF-8', standalone=True)
                break

        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)
        style.font.color.rgb = RGBColor(0, 0, 0)

        style.paragraph_format.space_after = Pt(6)

        heading_sizes = {1: 14, 2: 13, 3: 12, 4: 12}
        heading_space_before = {1: Pt(18), 2: Pt(14), 3: Pt(12), 4: Pt(10)}
        heading_space_after = {1: Pt(10), 2: Pt(8), 3: Pt(6), 4: Pt(6)}
        for level, size in heading_sizes.items():
            hs = doc.styles[f"Heading {level}"]
            hs.font.color.rgb = RGBColor(0, 0, 0)
            hs.font.name = "Times New Roman"
            hs.font.size = Pt(size)
            hs.font.bold = True
            hs.paragraph_format.space_before = heading_space_before[level]
            hs.paragraph_format.space_after = heading_space_after[level]

        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)

        return doc

    @classmethod
    def markdown_to_docx(
        cls,
        markdown_content: str,
        title: Optional[str] = None,
    ) -> bytes:
        """
        Convert Markdown content to DOCX format.

        Pipeline: clean markdown → HTML → DOCX (via htmldocx) with legal template.
        Uses htmldocx for high-fidelity conversion and a professional template
        with black headings, Calibri font, and proper margins.
        """
        import markdown as md
        from htmldocx import HtmlToDocx

        cleaned = re.sub(r"^\s*[-*_]{3,}\s*$", "", markdown_content, flags=re.MULTILINE)

        html = md.markdown(
            cleaned,
            extensions=["tables", "fenced_code", "sane_lists"],
        )

        html = re.sub(r"<hr\s*/?>", "", html)

        doc = cls._create_legal_template()
        converter = HtmlToDocx()
        converter.add_html_to_document(html, doc)

        cls._enforce_heading_spacing(doc)

        cls._enforce_font(doc)

        cls._fix_empty_list_items(doc)

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.read()

    @classmethod
    def _enforce_font(cls, doc, font_name: str = "Times New Roman"):
        """Force a consistent font on every run in the document.

        htmldocx-generated runs inherit Word theme fonts (Calibri) rather
        than the style-level font set in _create_legal_template().  Walking
        every run and setting font.name at the run level ensures the theme
        font is overridden.
        """
        for para in doc.paragraphs:
            for run in para.runs:
                run.font.name = font_name

    @classmethod
    def _enforce_heading_spacing(cls, doc):
        """Clear direct paragraph-format spacing on heading paragraphs.

        htmldocx sometimes sets explicit space_before/space_after on
        individual paragraphs, overriding the style's spacing.  Clearing
        the direct overrides lets the Heading style spacing (set in
        _create_legal_template) take effect.
        """
        from docx.oxml.ns import qn

        heading_styles = {"Heading 1", "Heading 2", "Heading 3", "Heading 4"}
        for para in doc.paragraphs:
            if para.style and para.style.name in heading_styles:

                pPr = para._p.find(qn("w:pPr"))
                if pPr is not None:
                    spacing = pPr.find(qn("w:spacing"))
                    if spacing is not None:
                        pPr.remove(spacing)

    @classmethod
    def _fix_empty_list_items(cls, doc):
        """Fix htmldocx bug: empty ListNumber/ListBullet paragraphs followed by content.

        htmldocx renders <ol><li>Text</li></ol> as two paragraphs:
          1) style=ListNumber, text=""   (empty — carries the numbering)
          2) style=Normal, text="Text"   (content — no numbering)

        This merges them: copies list pPr (style + numPr) to the content
        paragraph and removes the empty one.
        """
        from docx.oxml.ns import qn
        import copy

        LIST_STYLES = {
            "ListNumber", "ListBullet", "List Number", "List Bullet",
            "ListNumber2", "ListBullet2", "List Number 2", "List Bullet 2",
            "ListNumber3", "ListBullet3", "List Number 3", "List Bullet 3",
            "List Paragraph",
        }

        body = doc.element.body
        paragraphs = list(body.iterchildren(qn("w:p")))

        to_remove = []
        i = 0
        while i < len(paragraphs) - 1:
            p = paragraphs[i]
            p_next = paragraphs[i + 1]

            pPr = p.find(qn("w:pPr"))
            if pPr is None:
                i += 1
                continue

            pStyle = pPr.find(qn("w:pStyle"))
            if pStyle is None:
                i += 1
                continue

            style_name = pStyle.get(qn("w:val"))
            if style_name not in LIST_STYLES:
                i += 1
                continue

            text = "".join(
                t.text or ""
                for r in p.iterchildren(qn("w:r"))
                for t in r.iterchildren(qn("w:t"))
            )
            if text.strip():
                i += 1
                continue

            next_pPr = p_next.find(qn("w:pPr"))
            if next_pPr is not None:

                p_next.remove(next_pPr)
            p_next.insert(0, copy.deepcopy(pPr))

            to_remove.append(p)
            i += 2                        

        for p in to_remove:
            body.remove(p)

    @classmethod
    def get_docx_metadata(cls, docx_bytes: bytes) -> Dict[str, Any]:
        """
        Extract metadata from a DOCX file.

        Args:
            docx_bytes: Raw DOCX file bytes

        Returns:
            Dict with title, author, created, modified, etc.
        """
        Document = _ensure_docx()

        try:
            doc = Document(io.BytesIO(docx_bytes))
            props = doc.core_properties

            return {
                "title": props.title,
                "author": props.author,
                "created": props.created.isoformat() if props.created else None,
                "modified": props.modified.isoformat() if props.modified else None,
                "last_modified_by": props.last_modified_by,
                "revision": props.revision,
                "subject": props.subject,
                "keywords": props.keywords,
                "category": props.category,
                "comments": props.comments,
                "paragraph_count": len(doc.paragraphs),
                "word_count": sum(len(p.text.split()) for p in doc.paragraphs)
            }

        except Exception as e:
            logger.error(f"Failed to read DOCX metadata: {e}")
            return {}

def convert_docx_to_html(docx_bytes: bytes) -> str:
    """Convert DOCX bytes to HTML string."""
    html, _ = DocxService.docx_to_html(docx_bytes)
    return html

def convert_html_to_docx(html: str, title: Optional[str] = None) -> bytes:
    """Convert HTML string to DOCX bytes."""
    return DocxService.html_to_docx(html, title)

def convert_markdown_to_docx(markdown_content: str, title: Optional[str] = None) -> bytes:
    """Convert Markdown string to DOCX bytes."""
    return DocxService.markdown_to_docx(markdown_content, title)

def extract_text_from_docx(docx_bytes: bytes) -> str:
    """Extract plain text from DOCX bytes."""
    return DocxService.read_docx_text(docx_bytes)
